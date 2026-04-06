from __future__ import annotations

import io
from src.models import AssessmentResult, EvidenceItem

# ── Evidence tag labels ────────────────────────────────────────────────────────

_TAG_LABELS = {
    "directly_supported_by_input": "[supported]",
    "supported_by_input": "[supported]",
    "inferred_hypothesis": "[inferred]",
    "missing_evidence": "[missing]",
}

_TAG_COLORS_XLSX = {
    "directly_supported_by_input": "C6EFCE",   # light green
    "supported_by_input": "C6EFCE",
    "inferred_hypothesis": "FFEB9C",            # light yellow
    "missing_evidence": "FFC7CE",               # light red
}


def _tag_label(item: EvidenceItem) -> str:
    return _TAG_LABELS.get(item.evidence_tag, f"[{item.evidence_tag}]")


# ── Unicode sanitiser for PDF (Latin-1 safe) ──────────────────────────────────

_UNICODE_MAP: dict[str, str] = {
    "\u2014": "--",   # em dash
    "\u2013": "-",    # en dash
    "\u2022": "*",    # bullet
    "\u2019": "'",    # right single quotation mark
    "\u2018": "'",    # left single quotation mark
    "\u201c": '"',    # left double quotation mark
    "\u201d": '"',    # right double quotation mark
    "\u2026": "...",  # ellipsis
    "\u00a0": " ",    # non-breaking space
    "\u2192": "->",   # rightwards arrow
    "\u2190": "<-",   # leftwards arrow
    "\u2713": "OK",   # check mark
    "\u2715": "X",    # multiplication x
    "\u25b6": ">",    # black right-pointing triangle
    "\u00b0": "deg",  # degree sign
    "\u00d7": "x",    # multiplication sign
    "\u00e2": "a",    # a with circumflex (common encoding artifact)
}


def _safe(text: str) -> str:
    """Replace Unicode characters that Helvetica (Latin-1) cannot render."""
    for char, replacement in _UNICODE_MAP.items():
        text = text.replace(char, replacement)
    # Final pass: strip any remaining non-Latin-1 characters
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ── PDF export ─────────────────────────────────────────────────────────────────

def render_pdf_summary(result: AssessmentResult) -> bytes:
    """
    Generates a professional PDF report. Returns bytes for st.download_button.
    Requires fpdf2 (pip install fpdf2).
    """
    from fpdf import FPDF

    class _Report(FPDF):
        def header(self) -> None:
            self.set_font("Helvetica", "B", 8)
            self.set_text_color(130, 130, 130)
            self.cell(0, 7, _safe("LSS Copilot  |  Lean Six Sigma Assessment"), align="L")
            self.ln(10)

        def footer(self) -> None:
            self.set_y(-13)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(160, 160, 160)
            self.cell(0, 7, f"Page {self.page_no()}", align="C")

    pdf = _Report(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(left=20, top=16, right=20)
    pdf.add_page()

    # ── Title block ──
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(30, 27, 75)
    pdf.multi_cell(0, 11, _safe(result.project_name), align="L")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 120)
    pdf.cell(
        0, 7,
        _safe(f"Mode: {result.mode.upper()}   |   Audience: {result.audience.replace('_', ' ').title()}"),
        ln=True,
    )
    pdf.ln(2)
    pdf.set_draw_color(67, 97, 238)
    pdf.set_line_width(0.6)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(5)

    # ── Helpers ──
    def _section(title: str) -> None:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(67, 97, 238)
        pdf.cell(0, 7, _safe(title), ln=True)
        pdf.set_draw_color(200, 210, 255)
        pdf.set_line_width(0.3)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(3)
        pdf.set_text_color(33, 37, 41)

    def _subsection(title: str) -> None:
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(50, 50, 90)
        pdf.cell(0, 5, _safe(title), ln=True)
        pdf.set_text_color(33, 37, 41)

    def _body(text: str) -> None:
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(33, 37, 41)
        pdf.multi_cell(0, 6, _safe(text))

    def _bullet_items(items: list[EvidenceItem]) -> None:
        pdf.set_font("Helvetica", "", 9)
        for item in items:
            label = _tag_label(item)
            pdf.multi_cell(0, 5.5, _safe(f"  *  {item.statement}  {label}"))
            pdf.ln(0.8)

    def _bullet_strings(values: list[str]) -> None:
        pdf.set_font("Helvetica", "", 9)
        for v in values:
            pdf.multi_cell(0, 5.5, _safe(f"  *  {v}"))
            pdf.ln(0.8)

    def _table(rows: list[dict], col_widths: list[int], headers: list[str]) -> None:
        if not rows:
            return
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(240, 242, 255)
        pdf.set_text_color(30, 27, 75)
        for i, h in enumerate(headers):
            pdf.cell(col_widths[i], 7, _safe(h), border=1, fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(33, 37, 41)
        for row in rows:
            vals = [str(row.get(k, "")) for k in row]
            for i, val in enumerate(vals[:len(col_widths)]):
                pdf.cell(col_widths[i], 6, _safe(val[:80]), border=1)
            pdf.ln()

    # ── Content ──
    _section("Problem Statement")
    _body(result.cleaned_problem_statement)

    _section("Critical-to-Quality (CTQs)")
    _bullet_items(result.ctqs)

    _section("SIPOC")
    for col in ["suppliers", "inputs", "process", "outputs", "customers"]:
        _subsection(col.title())
        _bullet_strings(result.sipoc.get(col, []))

    _section("DMAIC Structure")
    for phase, items in result.dmaic_structure.items():
        _subsection(phase.upper())
        _bullet_items(items)

    _section("Root Causes")
    _bullet_items(result.root_causes)

    _section("Suggested Metrics")
    _bullet_items(result.suggested_metrics)

    _section("Improvement Actions")
    _bullet_items(result.improvement_actions)

    _section("Control Plan")
    _bullet_items(result.control_plan)

    _section("Action Tracker")
    _table(
        result.action_tracker,
        [82, 30, 26, 22],
        ["Action", "Owner", "Priority", "Status"],
    )

    _section("Project Memory")
    for key, values in result.project_memory.items():
        _subsection(key.replace("_", " ").title())
        _bullet_strings(values)

    _section("Role-Aware Summary")
    _body(result.role_summary)

    return bytes(pdf.output())


# ── DOCX export ────────────────────────────────────────────────────────────────

def render_docx_summary(result: AssessmentResult) -> bytes:
    """
    Generates a professional Word (.docx) document.
    Returns bytes for st.download_button.
    Requires python-docx (pip install python-docx).
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Helpers ──
    _BLUE = RGBColor(0x43, 0x61, 0xEE)
    _NAVY = RGBColor(0x1E, 0x1B, 0x4B)
    _GRAY = RGBColor(0x64, 0x74, 0x8B)

    _TAG_COLORS = {
        "directly_supported_by_input": RGBColor(0x15, 0x57, 0x24),
        "supported_by_input": RGBColor(0x15, 0x57, 0x24),
        "inferred_hypothesis": RGBColor(0x85, 0x64, 0x04),
        "missing_evidence": RGBColor(0x72, 0x1C, 0x24),
    }

    def _h1(text: str) -> None:
        p = doc.add_heading(text, level=1)
        p.runs[0].font.color.rgb = _NAVY
        p.runs[0].font.size = Pt(20)

    def _h2(text: str) -> None:
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            run.font.color.rgb = _BLUE
            run.font.size = Pt(13)

    def _h3(text: str) -> None:
        p = doc.add_heading(text, level=3)
        for run in p.runs:
            run.font.color.rgb = _NAVY
            run.font.size = Pt(11)

    def _body_para(text: str) -> None:
        doc.add_paragraph(text)

    def _evidence_bullet(item: EvidenceItem) -> None:
        para = doc.add_paragraph(style="List Bullet")
        run_text = para.add_run(item.statement + "  ")
        run_text.font.size = Pt(10)
        tag_color = _TAG_COLORS.get(item.evidence_tag, _GRAY)
        run_tag = para.add_run(_tag_label(item))
        run_tag.font.size = Pt(8)
        run_tag.font.color.rgb = tag_color
        run_tag.font.bold = True

    def _string_bullet(text: str) -> None:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(text)
        run.font.size = Pt(10)

    def _action_table(rows: list[dict]) -> None:
        if not rows:
            return
        headers = ["Action", "Owner", "Priority", "Status"]
        table = doc.add_table(rows=1 + len(rows), cols=4)
        table.style = "Light List Accent 5"
        # Header row
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
            hdr[i].paragraphs[0].runs[0].font.color.rgb = _NAVY
        # Data rows
        for r_idx, row in enumerate(rows):
            cells = table.rows[r_idx + 1].cells
            cells[0].text = str(row.get("action", ""))
            cells[1].text = str(row.get("owner", ""))
            cells[2].text = str(row.get("priority", ""))
            cells[3].text = str(row.get("status", ""))

    # ── Document content ──
    _h1(result.project_name)
    meta = doc.add_paragraph()
    meta.add_run(f"Mode: {result.mode.upper()}   |   Audience: {result.audience.replace('_', ' ').title()}")
    meta.runs[0].font.color.rgb = _GRAY
    meta.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    _h2("Problem Statement")
    _body_para(result.cleaned_problem_statement)

    _h2("Critical-to-Quality (CTQs)")
    for item in result.ctqs:
        _evidence_bullet(item)

    _h2("SIPOC")
    table = doc.add_table(rows=2, cols=5)
    table.style = "Light List Accent 5"
    cols = ["suppliers", "inputs", "process", "outputs", "customers"]
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c.title()
        hdr[i].paragraphs[0].runs[0].font.bold = True
    data = table.rows[1].cells
    for i, c in enumerate(cols):
        data[i].text = "\n".join(result.sipoc.get(c, []))

    _h2("DMAIC Structure")
    for phase, items in result.dmaic_structure.items():
        _h3(phase.upper())
        for item in items:
            _evidence_bullet(item)

    _h2("Root Causes")
    for item in result.root_causes:
        _evidence_bullet(item)

    _h2("Suggested Metrics")
    for item in result.suggested_metrics:
        _evidence_bullet(item)

    _h2("Improvement Actions")
    for item in result.improvement_actions:
        _evidence_bullet(item)

    _h2("Control Plan")
    for item in result.control_plan:
        _evidence_bullet(item)

    _h2("Action Tracker")
    _action_table(result.action_tracker)

    _h2("Project Memory")
    for key, values in result.project_memory.items():
        _h3(key.replace("_", " ").title())
        for v in values:
            _string_bullet(v)

    _h2("Role-Aware Summary")
    _body_para(result.role_summary)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX export ────────────────────────────────────────────────────────────────

def render_xlsx_summary(result: AssessmentResult) -> bytes:
    """
    Generates a structured Excel workbook with one sheet per section.
    Returns bytes for st.download_button.
    Requires openpyxl (pip install openpyxl).
    """
    import openpyxl
    from openpyxl.styles import (
        Font, PatternFill, Alignment, Border, Side
    )

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    # ── Style helpers ──
    _BLUE_FILL = PatternFill("solid", fgColor="4361EE")
    _GREEN_FILL = PatternFill("solid", fgColor="C6EFCE")
    _YELLOW_FILL = PatternFill("solid", fgColor="FFEB9C")
    _RED_FILL = PatternFill("solid", fgColor="FFC7CE")
    _BORDER_SIDE = Side(style="thin", color="D1D5DB")
    _THIN_BORDER = Border(
        left=_BORDER_SIDE, right=_BORDER_SIDE,
        top=_BORDER_SIDE, bottom=_BORDER_SIDE,
    )

    def _header_row(ws, row: int, values: list[str], fill=None) -> None:
        fill = fill or _BLUE_FILL
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.font = Font(bold=True, color="FFFFFF", size=10)
            cell.fill = fill
            cell.alignment = Alignment(wrap_text=True, vertical="center")
            cell.border = _THIN_BORDER

    def _data_row(ws, row: int, values: list[str], fill=None) -> None:
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = _THIN_BORDER
            if fill:
                cell.fill = fill

    def _evidence_fill(tag: str):
        if "supported" in tag:
            return _GREEN_FILL
        if "missing" in tag:
            return _RED_FILL
        return _YELLOW_FILL

    def _title_cell(ws, text: str) -> None:
        ws["A1"] = text
        ws["A1"].font = Font(bold=True, size=14, color="1E1B4B")
        ws["A2"] = f"Mode: {result.mode.upper()}   |   Audience: {result.audience.replace('_', ' ').title()}"
        ws["A2"].font = Font(size=10, color="64748B")

    def _set_col_widths(ws, widths: list[int]) -> None:
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    # ── Sheet 1: Overview ──
    ws1 = wb.create_sheet("Overview")
    _title_cell(ws1, result.project_name)
    ws1["A4"] = "Problem Statement"
    ws1["A4"].font = Font(bold=True, size=11, color="4361EE")
    ws1["A5"] = result.cleaned_problem_statement
    ws1["A5"].alignment = Alignment(wrap_text=True)

    ws1["A7"] = "Critical-to-Quality (CTQs)"
    ws1["A7"].font = Font(bold=True, size=11, color="4361EE")
    _header_row(ws1, 8, ["CTQ Statement", "Evidence"])
    for i, item in enumerate(result.ctqs, 9):
        _data_row(ws1, i, [item.statement, _tag_label(item)], _evidence_fill(item.evidence_tag))

    sipoc_start = 9 + len(result.ctqs) + 1
    ws1.cell(sipoc_start, 1, "SIPOC").font = Font(bold=True, size=11, color="4361EE")
    _header_row(ws1, sipoc_start + 1, ["Suppliers", "Inputs", "Process", "Outputs", "Customers"])
    max_rows = max(len(result.sipoc.get(c, [])) for c in ["suppliers", "inputs", "process", "outputs", "customers"])
    for r in range(max_rows):
        row_vals = [
            result.sipoc.get("suppliers", [""])[r] if r < len(result.sipoc.get("suppliers", [])) else "",
            result.sipoc.get("inputs", [""])[r] if r < len(result.sipoc.get("inputs", [])) else "",
            result.sipoc.get("process", [""])[r] if r < len(result.sipoc.get("process", [])) else "",
            result.sipoc.get("outputs", [""])[r] if r < len(result.sipoc.get("outputs", [])) else "",
            result.sipoc.get("customers", [""])[r] if r < len(result.sipoc.get("customers", [])) else "",
        ]
        _data_row(ws1, sipoc_start + 2 + r, row_vals)

    _set_col_widths(ws1, [50, 20, 20, 20, 20])

    # ── Sheet 2: DMAIC ──
    ws2 = wb.create_sheet("DMAIC")
    _title_cell(ws2, "DMAIC Structure")
    _header_row(ws2, 4, ["Phase", "Statement", "Evidence"])
    row = 5
    _PHASE_FILLS = {
        "define": PatternFill("solid", fgColor="DBEAFE"),
        "measure": PatternFill("solid", fgColor="EDE9FE"),
        "analyze": PatternFill("solid", fgColor="FEF3C7"),
        "improve": PatternFill("solid", fgColor="D1FAE5"),
        "control": PatternFill("solid", fgColor="CFFAFE"),
    }
    for phase, items in result.dmaic_structure.items():
        for item in items:
            _data_row(ws2, row, [phase.upper(), item.statement, _tag_label(item)], _PHASE_FILLS.get(phase))
            row += 1
    _set_col_widths(ws2, [12, 65, 14])

    # ── Sheet 3: Root Causes ──
    ws3 = wb.create_sheet("Root Causes")
    _title_cell(ws3, "Root Cause Analysis")
    _header_row(ws3, 4, ["Root Cause", "Evidence"])
    for i, item in enumerate(result.root_causes, 5):
        _data_row(ws3, i, [item.statement, _tag_label(item)], _evidence_fill(item.evidence_tag))
    _set_col_widths(ws3, [75, 15])

    # ── Sheet 4: Improvements ──
    ws4 = wb.create_sheet("Improvements")
    _title_cell(ws4, "Improvement Actions & Metrics")
    ws4["A4"] = "Improvement Actions"
    ws4["A4"].font = Font(bold=True, size=11, color="4361EE")
    _header_row(ws4, 5, ["Action", "Evidence"])
    for i, item in enumerate(result.improvement_actions, 6):
        _data_row(ws4, i, [item.statement, _tag_label(item)], _evidence_fill(item.evidence_tag))
    metrics_start = 6 + len(result.improvement_actions) + 1
    ws4.cell(metrics_start, 1, "Suggested Metrics").font = Font(bold=True, size=11, color="4361EE")
    _header_row(ws4, metrics_start + 1, ["Metric", "Evidence"])
    for i, item in enumerate(result.suggested_metrics, metrics_start + 2):
        _data_row(ws4, i, [item.statement, _tag_label(item)], _evidence_fill(item.evidence_tag))
    _set_col_widths(ws4, [75, 15])

    # ── Sheet 5: Control Plan ──
    ws5 = wb.create_sheet("Control Plan")
    _title_cell(ws5, "Control Plan")
    _header_row(ws5, 4, ["Control Item", "Evidence"])
    for i, item in enumerate(result.control_plan, 5):
        _data_row(ws5, i, [item.statement, _tag_label(item)], _evidence_fill(item.evidence_tag))
    at_start = 5 + len(result.control_plan) + 1
    ws5.cell(at_start, 1, "Action Tracker").font = Font(bold=True, size=11, color="4361EE")
    _header_row(ws5, at_start + 1, ["Action", "Owner", "Priority", "Status"])
    for i, row in enumerate(result.action_tracker, at_start + 2):
        _data_row(ws5, i, [
            row.get("action", ""), row.get("owner", ""),
            row.get("priority", ""), row.get("status", ""),
        ])
    _set_col_widths(ws5, [70, 18, 14, 14])

    # ── Sheet 6: Summary ──
    ws6 = wb.create_sheet("Summary")
    _title_cell(ws6, "Project Summary")
    ws6["A4"] = "Role-Aware Summary"
    ws6["A4"].font = Font(bold=True, size=11, color="4361EE")
    ws6["A5"] = result.role_summary
    ws6["A5"].alignment = Alignment(wrap_text=True)
    row = 7
    for key, values in result.project_memory.items():
        ws6.cell(row, 1, key.replace("_", " ").title()).font = Font(bold=True, color="1E1B4B")
        row += 1
        for v in values:
            ws6.cell(row, 1, f"  {v}")
            row += 1
        row += 1
    _set_col_widths(ws6, [90])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
